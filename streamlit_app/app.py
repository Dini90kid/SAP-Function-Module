# streamlit_app/app.py
import io
import os
import re
import sys
import textwrap
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

import pandas as pd
import streamlit as st

# ---------------------------------------------------------------------
# Ensure the repo root is importable for package imports
# ---------------------------------------------------------------------
BASE = Path(__file__).resolve().parent.parent
if str(BASE) not in sys.path:
    sys.path.insert(0, str(BASE))

# Use the same regex the runtime uses (keeps one source of truth)
from lakehouse_fm_agent.runtime.lineage import EDGE_RE, SKIP_RE  # type: ignore

APP_TITLE = "SAP FM Analyzer & Converter — Preview"
st.set_page_config(page_title="SAP FM Analyzer & Converter", layout="wide")
st.title(APP_TITLE)

# ---------------------------------------------------------------------
# Excel writer helper (auto-select engine)
# ---------------------------------------------------------------------
def workbook_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    engine = None
    try:
        import xlsxwriter  # noqa: F401
        engine = "xlsxwriter"
    except Exception:
        try:
            import openpyxl  # noqa: F401
            engine = "openpyxl"
        except Exception:
            st.error("Neither **XlsxWriter** nor **openpyxl** is available. "
                     "Add one of them to requirements.txt and redeploy.")
            return b""
    with pd.ExcelWriter(bio, engine=engine) as xw:
        for name, df in sheets.items():
            df.to_excel(xw, sheet_name=name[:31], index=False)
    bio.seek(0)
    st.caption(f"Excel writer engine used: **{engine}**")
    return bio.read()

# ---------------------------------------------------------------------
# Classification & analysis helpers
# ---------------------------------------------------------------------
STD_PREFIXES = ("R", "RS", "RR", "DD", "ENQUEUE", "DEQUEUE", "TR", "S", "CL_", "SAP")

def is_custom(fm: str) -> bool:
    if not isinstance(fm, str):
        return False
    return fm.upper().strip().startswith(("Y", "Z", "/"))

def category_for_fm(name: str) -> str:
    if not isinstance(name, str):
        return "Domain_Logic_Other"
    n = name.upper()
    if n.startswith(("RSD","RSR","RRSI","RSAU","RST","RSW","RSKC","RS_")): return "BW_Platform_API"
    if ("CURRENCY" in n) or ("CURR" in n) or ("UNIT_CONVERSION" in n) or ("UOM" in n) or ("BUOM" in n) or ("SSU" in n): return "UoM_Currency"
    if ("READ_0MATERIAL" in n) or ("READ_PRODFORM" in n) or ("READ_CUSTSALES" in n) or ("READ_ECLASS" in n) or ("READ_MASTER_DATA" in n): return "Masterdata_Readers"
    if ("DATE" in n) or ("MONTH" in n) or ("WEEK" in n) or ("PERIOD" in n) or n in {
        "SN_LAST_DAY_OF_MONTH","SLS_MISC_GET_LAST_DAY_OF_MONTH","/OSP/GET_DAYS_IN_MONTH","Y_PCM_LAST_DAY_OF_MONTH"
    }: return "Calendar_Time"
    if ("CONVERSION_EXIT_ALPHA" in n) or ("NUMERIC_CHECK" in n) or ("CHAVL_CHECK" in n) or ("REPLACE_STRANGE_CHARS" in n): return "Data_Cleansing_Validation"
    if "HIER" in n: return "Hierarchy"
    if "PAYTERMDAYS" in n: return "Payment_Terms"
    if "SID" in n or "TEXTS" in n: return "BW_SID_Texts"
    return "Domain_Logic_Other"

POINTERS = {
    "BW_Platform_API": "Replace BW mechanics with Lakehouse patterns: DSO → Delta MERGE; SID → joins by business keys; DDIC → control tables",
    "UoM_Currency": "Table‑driven conversion (FX, currency decimals, UoM factors) + Spark UDFs",
    "Masterdata_Readers": "Join curated dimensions in Unity Catalog; no generic BW readers",
    "Calendar_Time": "Use Spark date functions + enterprise calendar table",
    "Data_Cleansing_Validation": "Use Spark SQL expressions/UDFs (lpad, regex) instead of exits",
    "Hierarchy": "Parent‑child table + recursive/self‑joins",
    "Payment_Terms": "Model rules in a table; compute with PySpark",
    "BW_SID_Texts": "Keep attributes/texts in Delta and MERGE; no BW writers",
    "Domain_Logic_Other": "Assess & rebuild in PySpark with small reference tables",
}

def action_for_fm(fm: str) -> tuple[str, str]:
    cat = category_for_fm(fm)
    custom = is_custom(fm)
    if custom and cat in {"UoM_Currency","Masterdata_Readers","Hierarchy","Payment_Terms","Domain_Logic_Other"}:
        return "REBUILD", POINTERS.get(cat, POINTERS["Domain_Logic_Other"])
    if cat in {"BW_Platform_API","BW_SID_Texts"}:
        return "PATTERN", POINTERS[cat]
    if cat in {"UoM_Currency","Data_Cleansing_Validation","Calendar_Time"}:
        return "LIGHTWEIGHT", POINTERS[cat]
    if cat == "Masterdata_Readers":
        return "JOIN", POINTERS[cat]
    return ("REBUILD" if custom else "PATTERN"), POINTERS.get(cat, POINTERS["Domain_Logic_Other"])

def parse_lineage_text(text: str):
    edges = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        m_skip = SKIP_RE.search(line)
        if m_skip:
            p = m_skip.group("p").strip()
            c = m_skip.group("c").strip()
            edges.append((p, c))
            continue
        m = EDGE_RE.search(line.replace("  "," ").replace(" -","-"))
        if m:
            p = m.group("p").strip()
            c = m.group("c").strip()
            edges.append((p, c))
    # dedupe keeping order
    seen, uniq = set(), []
    for e in edges:
        if e not in seen:
            seen.add(e); uniq.append(e)
    return uniq

def edges_from_csv(df: pd.DataFrame):
    cols = {c: str(c).strip().lower().replace(" ", "_") for c in df.columns}
    df = df.rename(columns=cols)
    for need in ["main_fm","subfm","level2","level_3","level_4"]:
        if need not in df.columns:
            df[need] = None
    edges = []
    for _, r in df.iterrows():
        chain = [r.get("main_fm"), r.get("subfm"), r.get("level2"), r.get("level_3"), r.get("level_4")]
        chain = [x for x in chain if isinstance(x, str) and x.strip()!='']
        for a,b in zip(chain, chain[1:]):
            edges.append((a.strip(), b.strip()))
    # dedupe
    seen, uniq = set(), []
    for e in edges:
        if e not in seen:
            seen.add(e); uniq.append(e)
    return df, uniq

def degrees(edges):
    from collections import defaultdict
    in_deg, out_deg, nodes = defaultdict(int), defaultdict(int), set()
    for p,c in edges:
        nodes.add(p); nodes.add(c)
        out_deg[p]+=1; in_deg[c]+=1
        in_deg.setdefault(p, in_deg.get(p,0)); out_deg.setdefault(c, out_deg.get(c,0))
    return in_deg, out_deg, nodes

# Subgraph helpers
def build_adj(edges):
    from collections import defaultdict
    adj = defaultdict(set)
    for p,c in edges:
        adj[p].add(c)
    return adj

def bfs_nodes(adj, root, depth=3, direction="out"):
    seen, frontier, out_nodes = {root}, {root}, set()
    for _ in range(depth):
        nxt = set()
        for x in frontier:
            if direction in ("out","both"):
                for ch in adj.get(x, set()):
                    if ch not in seen:
                        seen.add(ch); nxt.add(ch); out_nodes.add(ch)
            if direction in ("in","both"):
                for pa, kids in adj.items():
                    if x in kids and pa not in seen:
                        seen.add(pa); nxt.add(pa); out_nodes.add(pa)
        frontier = nxt
        if not frontier:
            break
    return out_nodes

def build_mermaid_subgraph(adj, root, depth=3):
    nodes = bfs_nodes(adj, root, depth=depth, direction="out")
    edges = []
    for p, kids in adj.items():
        if p == root or p in nodes:
            for c in kids:
                if c in nodes or p == root:
                    edges.append((p,c))
    lines = ['graph TD'] + [f'  "{p}" --> "{c}"' for p,c in edges]
    return "```mermaid\n" + "\n".join(lines) + "\n```"

# ---------------------------------------------------------------------
# PROMPT parsing (refine list + depth)
# ---------------------------------------------------------------------
def parse_prompt_filters(prompt: str):
    """
    Supported directives (case-insensitive):
      category:<cat1>|<cat2>
      custom:true|false
      top:<N>
      depth:<D>
      include: FM1;FM2
      exclude: FM3;FM4
    """
    p = (prompt or "").strip()
    res = {
        "categories": None,
        "custom": None,          # True/False/None
        "top": None,             # int or None
        "depth": None,           # int or None
        "include": set(),
        "exclude": set(),
    }
    if not p:
        return res
    # category
    m = re.search(r"category\s*:\s*([^\n;]+)", p, flags=re.I)
    if m:
        res["categories"] = {x.strip() for x in re.split(r"[|,;]", m.group(1)) if x.strip()}
    # custom
    m = re.search(r"custom\s*:\s*(true|false)", p, flags=re.I)
    if m:
        res["custom"] = (m.group(1).lower() == "true")
    # top
    m = re.search(r"top\s*:\s*(\d+)", p, flags=re.I)
    if m:
        res["top"] = int(m.group(1))
    # depth
    m = re.search(r"depth\s*:\s*(\d+)", p, flags=re.I)
    if m:
        res["depth"] = int(m.group(1))
    # include
    m = re.search(r"include\s*:\s*([^\n]+)", p, flags=re.I)
    if m:
        res["include"] = {x.strip() for x in re.split(r"[;,]", m.group(1)) if x.strip()}
    # exclude
    m = re.search(r"exclude\s*:\s*([^\n]+)", p, flags=re.I)
    if m:
        res["exclude"] = {x.strip() for x in re.split(r"[;,]", m.group(1)) if x.strip()}
    return res

def apply_filters_to_nodes(nodes, catalog: pd.DataFrame, filters):
    items = list(nodes)
    df = catalog.copy()
    if filters["categories"]:
        df = df[df["Category"].isin(filters["categories"])]
    if filters["custom"] is not None:
        df = df[df["Is_Custom"] == ("Y" if filters["custom"] else "N")]
    # include/exclude override
    inc = filters["include"]
    exc = filters["exclude"]
    selected = set(df["FM"].tolist())
    if inc:
        selected |= {x for x in inc}
    if exc:
        selected -= {x for x in exc}
    # top N by criticality
    if filters["top"]:
        df2 = catalog[catalog["FM"].isin(selected)].sort_values(
            ["Criticality","In_Degree","Out_Degree"], ascending=False
        ).head(filters["top"])
        selected = set(df2["FM"].tolist())
    # keep original order but filter to selected
    return [n for n in items if n in selected] if selected else items

# ---------------------------------------------------------------------
# Build MASTER.docx
# ---------------------------------------------------------------------
def build_master_docx_bytes(fm: str,
                            category: str,
                            action: str,
                            rationale: str,
                            indeg: int,
                            outdeg: int,
                            subgraph_mermaid: str,
                            prompt: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(f"{fm} — Master Document", level=0)

    p = doc.add_paragraph()
    run = p.add_run(f"Category: {category}   |   Action: {action}   |   in_degree: {indeg}   |   out_degree: {outdeg}")
    run.bold = True

    doc.add_heading("1. What this FM does (AS‑IS)", level=1)
    doc.add_paragraph(
        "Inferred from lineage and category mapping. This FM is categorized as "
        f"{category}; follow pattern guidance rather than 1:1 port."
    )

    doc.add_heading("2. Lineage (focused Mermaid, paste in docs)", level=1)
    code = doc.add_paragraph()
    run = code.add_run(subgraph_mermaid)
    font = run.font; font.name = "Courier New"; font.size = Pt(10)

    doc.add_heading("3. How to approach (TO‑BE)", level=1)
    doc.add_paragraph(f"Selected action: {action}")
    doc.add_paragraph(f"Rationale: {rationale}")
    doc.add_paragraph(
        "- Replace BW mechanics with Lakehouse patterns; use table‑driven conversion; "
        "Spark SQL/UDFs; enterprise calendar; business‑key joins (no SIDs)."
    )

    doc.add_heading("4. Converted code (skeleton)", level=1)
    doc.add_paragraph(
        "See handler.py in the downloaded ZIP for a runnable skeleton tailored for this FM. "
        "Augment with real tables, joins, and UDFs based on project contracts."
    )

    doc.add_heading("5. Handling nested FMs", level=1)
    doc.add_paragraph(
        "Nested FMs (within chosen depth) are summarized in nested_action_plan.csv (included in ZIP): "
        "REBUILD (custom), PATTERN (BW platform), LIGHTWEIGHT (SQL/UDF), JOIN (masterdata readers)."
    )

    doc.add_heading("6. Prompt notes", level=1)
    doc.add_paragraph(prompt.strip() or "—")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ---------------------------------------------------------------------
# Upload: ZIP/TXT/CSV/XLSX
# ---------------------------------------------------------------------
uploaded = st.file_uploader(
    "Upload ZIP (multiple run folders with `LINEAGE.txt`), or a single LINEAGE.txt, or your CSV/XLSX",
    type=["zip","txt","csv","xlsx"],
)

if not uploaded:
    st.info("Upload a **.zip** / **.txt** / **.csv** / **.xlsx** to view the analysis, prompt, and docs.")
    st.stop()

runs_edges = {}
your_format_df = None
zip_lineage_paths = []

if uploaded.name.lower().endswith(".zip"):
    z = ZipFile(io.BytesIO(uploaded.getvalue()))
    lineage_members = [
        n for n in z.namelist()
        if (n.upper().endswith("LINEAGE.TXT") or ("LINEAGE" in n.upper() and n.upper().endswith(".TXT")))
    ]
    zip_lineage_paths = lineage_members[:]
    if not lineage_members:
        st.error("No `LINEAGE*.txt` found in the ZIP. Ensure each run folder contains one.")
        st.stop()
    for member in lineage_members:
        try:
            text = z.read(member).decode("utf-8", errors="ignore")
        except Exception:
            text = z.read(member).decode("latin-1", errors="ignore")
        parts = [p for p in member.split("/") if p]
        run_key = parts[0] if parts else member
        edges = parse_lineage_text(text)
        if edges:
            runs_edges.setdefault(run_key, []).extend(edges)

elif uploaded.name.lower().endswith(".txt"):
    text = uploaded.getvalue().decode("utf-8", errors="ignore")
    runs_edges["SINGLE_RUN"] = parse_lineage_text(text)

elif uploaded.name.lower().endswith(".csv"):
    df_in = pd.read_csv(uploaded)
    your_format_df, edges = edges_from_csv(df_in)
    runs_edges["FROM_CSV"] = edges

else:  # .xlsx
    try:
        x = pd.ExcelFile(uploaded, engine="openpyxl")
    except Exception:
        x = pd.ExcelFile(uploaded)  # let pandas choose
    if "Interdependency_Edges" in x.sheet_names:
        ed = x.parse("Interdependency_Edges")
        if {"Parent","Child"}.issubset(ed.columns):
            edges = list(zip(ed["Parent"].astype(str), ed["Child"].astype(str)))
            runs_edges["FROM_XLSX"] = edges
    try:
        your_format_df = x.parse("Your_Format")
    except Exception:
        your_format_df = x.parse(x.sheet_names[0])

if zip_lineage_paths:
    st.caption(f"Found LINEAGE files in ZIP: {len(zip_lineage_paths)}")
    with st.expander("Show first 10 LINEAGE paths found"):
        st.code("\n".join(zip_lineage_paths[:10]))

# Combine & de-dupe edges
combined_edges = []
for e in runs_edges.values():
    combined_edges.extend(e)
seen, uniq = set(), []
for e in combined_edges:
    if e not in seen:
        seen.add(e); uniq.append(e)
combined_edges = uniq

if not combined_edges:
    st.error("No edges parsed. ZIP must contain lineage files with lines like `PARENT -> CHILD [ FM ]`, "
             "or upload your CSV/XLSX with hierarchy columns.")
    st.stop()

# ---------------------------------------------------------------------
# Build analysis tables (Your Format + rich views)
# ---------------------------------------------------------------------
in_deg, out_deg, nodes = degrees(combined_edges)
adj = build_adj(combined_edges)

def make_your_format(df_source: pd.DataFrame | None, edges: list[tuple[str,str]]):
    cols = ["Main FM","SubFM","Level2","Level 3","Level 4","Type","Remarks","Action Plan","Download Document"]
    if df_source is not None:
        df = df_source.copy()
        for c in cols:
            if c not in df.columns: df[c] = ""
        order = ["Level 4","Level 3","Level2","SubFM","Main FM"]
        def best_name(row):
            for pos in order:
                val = row.get(pos, "")
                if isinstance(val, str) and val.strip():
                    return val.strip()
            return ""
        for i,row in df.iterrows():
            name = best_name(row)
            if name:
                act, why = action_for_fm(name)
                if not str(df.at[i,"Type"]).strip(): df.at[i,"Type"] = "Custom" if is_custom(name) else "Standard"
                if not str(df.at[i,"Action Plan"]).strip(): df.at[i,"Action Plan"] = act
                if not str(df.at[i,"Remarks"]).strip(): df.at[i,"Remarks"] = why
        return df[cols]
    else:
        from collections import defaultdict
        kids = defaultdict(set)
        for p,c in edges: kids[p].add(c)
        rows = []
        for p in sorted(kids.keys()):
            for c in sorted(kids[p]):
                gks = sorted(kids.get(c, []))
                if gks:
                    for g in gks:
                        rows.append([p, c, g, "", "", "Custom" if is_custom(g) else "Standard",
                                     POINTERS.get(category_for_fm(g), ""),
                                     action_for_fm(g)[0], ""])
                else:
                    rows.append([p, c, "", "", "", "Custom" if is_custom(c) else "Standard",
                                 POINTERS.get(category_for_fm(c), ""),
                                 action_for_fm(c)[0], ""])
        return pd.DataFrame(rows, columns=cols)

your_format = make_your_format(your_format_df, combined_edges)

cat_rows = []
for fm in sorted(nodes):
    act, why = action_for_fm(fm)
    cat_rows.append({
        "FM": fm,
        "Category":  category_for_fm(fm),
        "Is_Custom": "Y" if is_custom(fm) else "N",
        "Action":     act,
        "Rationale":  why,
        "In_Degree":  in_deg.get(fm,0),
        "Out_Degree": out_deg.get(fm,0),
        "Criticality": in_deg.get(fm,0)*2 + out_deg.get(fm,0) + (1 if is_custom(fm) else 0),
    })
catalog_cols = ["FM","Category","Is_Custom","Action","Rationale","In_Degree","Out_Degree","Criticality"]
catalog = pd.DataFrame(cat_rows, columns=catalog_cols)
if len(catalog)>0:
    catalog = catalog.sort_values(["Criticality","In_Degree","Out_Degree"], ascending=False)

if your_format_df is not None and "Main FM" in your_format_df.columns:
    mains = [m for m in your_format_df["Main FM"].dropna().unique().tolist() if isinstance(m, str) and m.strip()]
else:
    mains = sorted(set(p for p,_ in combined_edges))

summary_rows = []
for main in mains:
    direct = sorted(adj.get(main, []))
    # nested (≤3 hops)
    seen2, frontier, nested = {main}, {main}, set()
    for _ in range(3):
        nxt = set()
        for x in frontier:
            for ch in adj.get(x, set()):
                if ch not in seen2:
                    seen2.add(ch); nxt.add(ch); nested.add(ch)
        frontier = nxt
        if not frontier:
            break
    create_list  = sorted(x for x in nested if is_custom(x))
    pattern_list = sorted(x for x in nested if not is_custom(x))
    summary_rows.append({
        "Main_FM": main,
        "Direct_Calls_Count": len(direct),
        "Nested_Upto3_Count": len(nested),
        "Create_Handlers_Count": len(create_list),
        "Pattern_Replace_Count": len(pattern_list),
        "Create_Handlers_List": ", ".join(create_list),
        "Pattern_Replace_List": ", ".join(pattern_list),
    })
main_summary = pd.DataFrame(summary_rows)

overall_build = (catalog.loc[catalog.get("Is_Custom","N")=="Y", ["FM","Category","Action","Criticality"]]
                 if len(catalog)>0 else pd.DataFrame(columns=["FM","Category","Action","Criticality"]))
if len(overall_build)>0:
    overall_build = overall_build.sort_values(["Action","Criticality"], ascending=[True, False])

edges_df = pd.DataFrame(combined_edges, columns=["Parent","Child"])
edges_df["Parent_Category"] = edges_df["Parent"].apply(category_for_fm)
edges_df["Child_Category"]  = edges_df["Child"].apply(category_for_fm)
edges_df["Child_Is_Custom"] = edges_df["Child"].apply(lambda x: "Y" if is_custom(x) else "N")

matrix = pd.crosstab(edges_df["Parent"], edges_df["Child"]) if len(edges_df)>0 else pd.DataFrame()

# ---------------------------------------------------------------------
# Prompt input + PROCESS button
# ---------------------------------------------------------------------
st.subheader("Give a prompt for the next step")
default_prompt = ("Generate Master/LLD/Test docs and handler.py for the top:5 custom:true category:UoM_Currency|Calendar_Time "
                  "depth:3; include:Y_DNP_CONV_BUOM_SU_SSU; exclude:RSD_CHAVL_READ_ALL")
user_prompt = st.text_area("Your prompt", value=default_prompt, height=120)

filters = parse_prompt_filters(user_prompt)
if filters.get("depth"):
    doc_depth_default = max(1, min(6, filters["depth"]))
else:
    doc_depth_default = 3

if st.button("Process prompt"):
    st.session_state["filters"] = filters
    st.success("Prompt processed. The FM list below is now refined.")
else:
    if "filters" not in st.session_state:
        st.session_state["filters"] = filters

# ---------------------------------------------------------------------
# On-screen display (tabs)
# ---------------------------------------------------------------------
tabs = st.tabs([
    "Your Format", "FM Catalog / Action Plan", "Per Main Summary",
    "Edges", "Matrix", "Overall Build List"
])

with tabs[0]:
    st.dataframe(your_format, use_container_width=True, hide_index=True)
with tabs[1]:
    st.dataframe(catalog, use_container_width=True, hide_index=True)
with tabs[2]:
    st.dataframe(main_summary, use_container_width=True, hide_index=True)
with tabs[3]:
    st.dataframe(edges_df, use_container_width=True, hide_index=True)
with tabs[4]:
    st.dataframe(matrix, use_container_width=True)
with tabs[5]:
    st.dataframe(overall_build, use_container_width=True, hide_index=True)

# Excel download (includes Prompt_Box)
wb_sheets = {
    "Your_Format": your_format,
    "FM_Catalog_ActionPlan": catalog,
    "Per_Main_Summary": main_summary,
    "Interdependency_Edges": edges_df,
    "Interdependency_Matrix": matrix.reset_index() if len(matrix)>0 else pd.DataFrame(),
    "Overall_Build_List": overall_build,
    "Prompt_Box": pd.DataFrame({"Your Prompt Here:":[user_prompt]}),
}
wb_bytes = workbook_bytes(wb_sheets)
st.download_button(
    "⬇️ Download full analysis workbook (Excel)",
    data=wb_bytes,
    file_name="SAP_FM_Analysis.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# ---------------------------------------------------------------------
# Design & Docs Generator
# ---------------------------------------------------------------------
st.subheader("Design & Docs Generator (per FM)")

# Apply prompt-based filters to nodes for the FM selector
nodes_ordered = sorted(nodes)
nodes_filtered = apply_filters_to_nodes(nodes_ordered, catalog, st.session_state["filters"])

fm_pick = st.selectbox("Choose FM", options=["<none>"] + nodes_filtered)
doc_depth = st.slider("Lineage depth for docs", min_value=1, max_value=6, value=doc_depth_default)

st.caption("Generates: Master.md + LLD.md + Test.md + handler.py + lineage.mmd + nested_action_plan.csv (in a ZIP). "
           "Also a separate Word version of MASTER.")

if fm_pick != "<none>":
    # Build doc pack (ZIP)
    files_dict = {}
    # Compute metadata
    row = catalog.loc[catalog["FM"]==fm_pick]
    category = row["Category"].iloc[0] if len(row)>0 else category_for_fm(fm_pick)
    action, why = action_for_fm(fm_pick) if len(row)==0 else (row["Action"].iloc[0], row["Rationale"].iloc[0])
    indeg = in_deg.get(fm_pick,0)
    outdeg = out_deg.get(fm_pick,0)
    sub_mermaid = build_mermaid_subgraph(adj, fm_pick, depth=doc_depth)

    # Build ZIP contents
    def make_handler_py(fm: str, category: str, prompt: str) -> str:
        return f'''# handler.py — generated {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
# FM: {fm}
# Category: {category}
# Prompt notes: {prompt.strip() or "—"}

from lakehouse_fm_agent.runtime.context import Context
from typing import Optional

def handler(ctx: Optional[Context]) -> None:
    """
    PURPOSE (AS-IS):
      - Implement logic formerly in FM {fm} in a Lakehouse-native way.

    TO-BE (Approach):
      - Use category-driven patterns:
        * BW Platform  -> MERGE/joins/control tables (no SID, no buffer)
        * UoM/Currency -> table-driven conversion + UDFs (DECIMAL-safe)
        * Masterdata   -> joins to curated dimensions
        * Calendar     -> Spark date functions + enterprise calendar
        * Cleansing    -> Spark SQL expressions/UDFs

    Implementation sketch:
      1) Read inputs (ctx.current_df or ctx.read_uc(...))
      2) Apply transforms (joins/UDFs/SQL) per rules & prompt guidance
      3) Write/merge outputs or set ctx.current_df

    Notes from prompt:
      {textwrap.indent(prompt.strip() or "—", "      ")}
    """
    # TODO: implement with real contracts
    pass
'''
    # nested action plan
    sub_nodes = bfs_nodes(adj, fm_pick, depth=doc_depth, direction="out")
    rows_nested = []
    for n in sorted(sub_nodes):
        aa, ww = action_for_fm(n)
        rows_nested.append([n, category_for_fm(n), "Y" if is_custom(n) else "N", aa, ww])
    nested_df = pd.DataFrame(rows_nested, columns=["FM","Category","Is_Custom","Action","Rationale"])
    nested_csv = nested_df.to_csv(index=False).encode("utf-8")

    master_md = f"""# {fm_pick} — Master Document

**Category:** {category}  |  **Action:** **{action}**  |  **in_degree:** {indeg}  |  **out_degree:** {outdeg}

## 1. What this FM does (AS‑IS)
- Inferred from lineage and category mapping.

## 2. Lineage (focused view: depth={doc_depth})
{sub_mermaid}

## 3. How to approach (TO‑BE)
- Replace BW mechanics with Lakehouse patterns; table‑driven conversion; Spark SQL/UDFs; enterprise calendar; business‑key joins (no SIDs).
- Selected action: **{action}**
- Rationale: {why}

## 4. Converted code (skeleton)
See `handler.py` in this ZIP. Augment with contracts (tables, keys, decimals).

## 5. Handling nested FMs
- See `nested_action_plan.csv` (REBUILD / PATTERN / LIGHTWEIGHT / JOIN)

## 6. Prompt notes
{user_prompt.strip() or "—"}
"""

    lld_md = f"""# {fm_pick} — Design LLD

## Components
- Handler: `handler.py`
- Category: **{category}**
- Data contracts: define schemas & DECIMAL scales for money/UoM

## Execution
1) Read reference tables (FX, currency_decimals, uom_factors, calendar, logsys_map) as needed
2) Apply transformations (joins/UDFs/SQL) according to Action **{action}**
3) MERGE/WRITE outputs or return DataFrame via context

## Error Handling & Observability
- Null/invalid checks, row counts, rejects
- Log applied patterns and nested decisions

## Prompt notes
{user_prompt.strip() or "—"}
"""

    test_md = f"""# {fm_pick} — How to Test

## Unit tests
- Golden cases for UoM/currency conversion (scale/rounding)
- Cleansing (alpha/numeric)

## Integration tests
- Replay a short path around **{fm_pick}**
- Validate row counts, key uniqueness, and business rules
"""

    files_dict = {
        f"{fm_pick}_MASTER.md": master_md.encode("utf-8"),
        f"{fm_pick}_LLD.md": lld_md.encode("utf-8"),
        f"{fm_pick}_TEST.md": test_md.encode("utf-8"),
        "handler.py": make_handler_py(fm_pick, category, user_prompt).encode("utf-8"),
        "lineage.mmd": sub_mermaid.encode("utf-8"),
        "nested_action_plan.csv": nested_csv,
        "ABOUT.txt": f"Generated: {datetime.now().isoformat()}\nFM: {fm_pick}\nDepth: {doc_depth}\n".encode("utf-8"),
    }
    zbytes = io.BytesIO()
    from zipfile import ZipFile, ZIP_DEFLATED
    with ZipFile(zbytes, "w", ZIP_DEFLATED) as zf:
        for name, data in files_dict.items():
            zf.writestr(name, data)
    zbytes.seek(0)

    # Build a Word (.docx) version of MASTER
    master_docx_bytes = build_master_docx_bytes(
        fm=fm_pick,
        category=category,
        action=action,
        rationale=why,
        indeg=indeg,
        outdeg=outdeg,
        subgraph_mermaid=sub_mermaid,
        prompt=user_prompt,
    )

    # Two download buttons
    st.download_button(
        "⬇️ Download Doc Pack (ZIP)",
        data=zbytes.getvalue(),
        file_name=f"{fm_pick}_Doc_Pack.zip",
        mime="application/zip"
    )
    st.download_button(
        "⬇️ Download MASTER (Word)",
        data=master_docx_bytes,
        file_name=f"{fm_pick}_MASTER.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info("Select an FM above, adjust depth/prompt as needed, then download ZIP and/or Word doc.")
